# 08_docx_in_datei_schreiben.py
# In diesem Beispielscript geht es darum einen Beliebigen String in ein Word-Dokument zu 08_docx_in_datei_schreiben

import os, docx
from docx.enum.text import WD_BREAK_TYPE
os.chdir(os.path.dirname(__file__))

target_file='.\\hello_world.docx'
if os.path.exists(target_file):
    os.remove(target_file)

print('Bitte gewünschten String eingeben')
user_paragraph=input('String: ')
doc_file=docx.Document()

# Mit add_paragraph() kann man dem Dokument neue Paragraphen hinzufügen
doc_string1=doc_file.add_paragraph(user_paragraph).style='Title'
doc_string2=doc_file.add_paragraph('Das ist ein weiterer Paragraph. ').style='Heading 1'

# Der Style kann auch als 2. Argument bei add_paragraph() übergeben werden.
doc_string3=doc_file.add_paragraph('Und das ist auch ein Paragraph. ', 'Quote')

# Mit add_run() lassen sich neue Runs zu einem Paragraphen hinzufügen
run_object=doc_string3.add_run('Dieser Text ist ein weiterer run welcher dem Paragraph 3 hinzugefügt wird.').underline=True

# Mit add_heading() kann man neue Headings hinzufügen
doc_file.add_paragraph('\n')
doc_file.add_heading('Header 0', 0)
doc_file.add_heading('Header 1', 1)
doc_file.add_heading('Header 2', 2)
doc_file.add_heading('Header 3', 3)
doc_file.add_heading('Header 4', 4)

# Um eine neue Seite zu beginnen verwendet man doc_file.paragraphs[x].runs[y].add_break(WD_BREAK_TYPE.PAGE). Dazu benötigt man die Funktion WD_BREAK_TYPE aus docx.enum.text.
doc_file.paragraphs[8].runs[0].add_break(WD_BREAK_TYPE.PAGE)
doc_file.add_paragraph('Dies ist eine neue Seite', 'Heading 2')

# Mit add.picture() lassen sich Bilder einfügen. Die Grösse des Bildes muss entweder in Inches oder in cm angegeben werden mit docx.shared.inches(1) oder docx.shared.Cm(4). Wenn man die Definition der Grösse auslässt wird das Bild in seiner Originalgrösse eingefügt.
doc_file.add_paragraph('Dies ist ein Bild einer Katze:', 'Normal')
doc_file.add_picture('zophie.png', width=docx.shared.Mm(50), height=docx.shared.Cm(8))

doc_file.save(target_file)

