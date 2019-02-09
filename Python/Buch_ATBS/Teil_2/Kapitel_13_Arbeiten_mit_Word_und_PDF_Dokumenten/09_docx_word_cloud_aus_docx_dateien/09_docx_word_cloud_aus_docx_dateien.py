# 09_docx_word_cloud_aus_docx_dateien.py
# In dieser Übung geht es darum den Kompletten Text aus einem Word-Dokument zu erfassen und
# die darin vorkommenden Wörter zu zählen und anschliessend als Word-Cloud auszugeben bei
# welcher die Schriftgrösse des Wortes dessen vorkommen repräsentiert.
# Prüfen ob Nomen oder nicht: http://www.learnersdictionary.com/definition/<WORD>

import os, docx, re
from docx.shared import Pt
from docx.shared import RGBColor

os.chdir(os.path.dirname(__file__))

target_file='.\\word_cloud.docx'
source_file='.\\read_me.docx'

while True:
    try:
        if os.path.exists(target_file):
            os.remove(target_file)
        break
    except:
        print('File ist gesperrt da noch geöffnet. Bitte word_cloud.docx schliessen.')
        input('Enter zum Fortfahren')

# Mindestlänge eines Wortes damit es Gezählt wird
minimum_word_length=1
# Minimales vorkommen des Wortes um aufgeführt zu werden
minimum_word_amount=5
# Grösste Schriftart, alle Wörter die mehr als cut_off vorkommen werden 
# Rot markiert und auf Schriftgrösse 80 belassen.
cut_off=80

# Zu ignorierende Wörter.
ignore_list=['is','the','if','then','that','and', 'i', 'not', 'had', 'for', 'was']
# ignore_list=[]

# Durchsuche Datei nach Wörtern und Zähle wie oft diese vorkommen.
word_counter={}
read_me_file=docx.Document(source_file)
for paragraph in read_me_file.paragraphs:
    for run in paragraph.runs:
        search_pattern=re.compile(r'\W?(\w+)\W?')
        words_in_run=search_pattern.findall(run.text)
        for word in words_in_run:
            # Wenn das Wort länger als das Minimum ist speichern
            if len(word) >= minimum_word_length:
                word_counter.setdefault(word.lower(), 0)
                word_counter[word.lower()]+=1

# Erstellen des Dokuments
docx_file=docx.Document()
docx_file.add_paragraph('Wordcloud aus der Datei '+source_file, 'Heading 1')
for i in range(2):
    docx_file.add_paragraph('')

# Wörter dem Dokument hinzufügen
hit_cut_off=False
for word in word_counter:
    if word_counter[word] >= minimum_word_amount:
        if word not in ignore_list:
            run=docx_file.paragraphs[len(docx_file.paragraphs)-1].add_run(word+' ')
            if word_counter[word] > cut_off:
            # Wenn das Wort mehr als der Wert cut_off mal vorkommt setze Grösse auf den Wert cut_off und markiere rot
                run.italic=True
                font=run.font
                font.size=Pt(cut_off)
                font.color.rgb = RGBColor(0x88, 0x00, 0x00)
                hit_cut_off=True
            else:
            # Setze Schrifgrösse auf anzahl des entsprechenden Wortes
                font=run.font
                font.size=Pt(word_counter[word])
if hit_cut_off:
    for i in range(2):
        docx_file.add_paragraph('')
    run=docx_file.paragraphs[len(docx_file.paragraphs)-1].add_run('Text in dieser Schrift und Farbe kommt mehr als '+str(cut_off)+' mal im Text vor.')
    run.italic=True
    font=run.font
    font.size=Pt(14)
    font.color.rgb = RGBColor(0x88, 0x00, 0x00)
docx_file.save(target_file)